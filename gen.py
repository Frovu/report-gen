#!/usr/bin/python

import imgkit
import os
import sys

def clean():
	if os.path.exists('img'):
		for f in os.listdir('img'):
			os.remove(os.path.join('img', f))

def gen_image(path):
	files = os.listdir(path)
	css = [f for f in files if f.endswith('.css')]
	fp = os.path.join(path, 'index.html')
	try:
		imgkit.from_file(fp, os.path.join('img', f'{path.split("/")[-1]}.jpg'),
			options={'quiet': '', 'enable-local-file-access': ''}, css=css)
	except Exception as e:
		print(f'exception in rendering {path.split("/")[-1]} html \n{e}')

def read_tasks(path):
	task = []
	with open(path) as f:
		for part in f.read().split('<h3>')[1:]:
			task.append([from_html(a.strip()) for a in part.split('</h3>')])
	return task

text_to_html = {
	'&': ['&amp;'],
	'<': ['&lt;', '&lt'],
	'>': ['&gt;', '&gt'],
	'\n': ['<br>']
}
def from_html(text):
	newtext = text
	for i in text_to_html:
		for j in text_to_html[i]:
			newtext = newtext.replace(j, i)
	return newtext

def build_all(path, overwrite=False):
	if not os.path.exists('img'):
		os.makedirs('img')
	codes = []
	images = [f for f in os.listdir('img') if f.endswith('.jpg')]
	for (dirpath, dirnames, filenames) in os.walk(path):
		dirnames.sort(key=lambda x: int(x.split('task')[1]))
		for d in dirnames:
			dp = os.path.join(dirpath, d)
			print(f'building {d}...', end='', flush=True)
			if overwrite or f'{d}.jpg' not in images:
				gen_image(dp)
			code = []
			for f in os.listdir(dp):
				if f.split('.')[-1] in ['html', 'js', 'css']:
					with open(os.path.join(dirpath, d, f)) as file:
						code.append((f, file.read()))
			codes.append(code)
			print('done')
		break
	return codes

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

document = Document('template.docx')
heading = document.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH, True)
heading.font.name = 'Times New Roman'
heading.font.bold = True
heading.font.size = Pt(14)
code_style = document.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Courier New'
code_style.font.size = Pt(10)
def_style = document.styles.add_style('Default', WD_STYLE_TYPE.PARAGRAPH)
def_style.font.name = 'Times New Roman'
def_style.font.size = Pt(14)
#def_style.paragraph_format.left_indent = Cm(1)
def_style.paragraph_format.line_spacing = Pt(24)
def_style.paragraph_format.space_after = Pt(12)

def build_report(path):
	codes = build_all(path, len(sys.argv) > 2 and sys.argv[2] == 'rebuild')
	images = os.listdir('img')
	tasks = read_tasks(os.path.join(path, 'task.html'))
	print(f'images amount = {len(images)}')
	print(f'tasks amount = {len(tasks)}')
	print('generatig document..')

	for i, task in enumerate(tasks):
		document.add_heading(task[0], 2)
		document.add_paragraph()
		p = document.add_paragraph(style='Default')
		p.paragraph_format.first_line_indent = Cm(1.25)
		p.add_run('Задание:').bold = True
		p = document.add_paragraph(style='Default')
		p.paragraph_format.first_line_indent = Cm(1.25)
		p.add_run(task[1])
		p = document.add_paragraph(style='Default')
		p.add_run('Описание выполнения задания').bold = True
		p = document.add_paragraph('TODO', style='List Paragraph')
		p.paragraph_format.line_spacing = Pt(24)
		p.style.font.size = Pt(14)
		if i < len(codes):
			for ci, c in enumerate(codes[i]):
				document.add_paragraph(f'Листинг {i+1}.{ci+1} – {c[0]}', style='Default')
				document.add_paragraph(c[1], style='Code')
		image = next((a for a in images if f"task{str(i+1)}." in a), None)
		if image:
			ipr = document.add_paragraph().add_run()
			ipr.add_picture(os.path.join("img", image), width=Cm(16))
		p = document.add_paragraph(f'Рисунок {i+1} - Результат выполнения кода на странице')
		p.style = 'Default'
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		else:
			print(f'image not found for task {str(i+1)}')
		document.add_page_break()
	print('done!')
	document.save('output.docx')

if len(sys.argv) == 1:
	print(f'Usage: {sys.argv[0]} <path>|clean [rebuild]')
elif sys.argv[1] == 'clean':
	clean()
else:
	build_report(sys.argv[1])
